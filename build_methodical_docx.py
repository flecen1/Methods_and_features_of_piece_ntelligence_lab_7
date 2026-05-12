# -*- coding: utf-8 -*-
"""
Генерація методичного документа Word (.docx) до практичної роботи з ГА на Python.
Запуск: python build_methodical_docx.py
Вихід: Методичні_вказівки_Практична_робота_ГА_Python.docx (+ ASCII-копія).
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

GITHUB_REPO_WEB = "https://github.com/flecen1/Methods_and_features_of_piece_ntelligence_lab_7"
GITHUB_CLONE_HTTPS = "https://github.com/flecen1/Methods_and_features_of_piece_ntelligence_lab_7.git"
REPO_FOLDER_NAME = "Methods_and_features_of_piece_ntelligence_lab_7"


def _set_normal_font(doc: Document, name: str = "Times New Roman", size: int = 14) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def build(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    _set_normal_font(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(
        "МЕТОДИЧНІ ВКАЗІВКИ\n"
        "до практичної роботи з дисципліни\n"
        "«Моделі та системи штучного інтелекту»\n\n"
        "Практична робота\n"
        "«Дослідження генетичних алгоритмів на задачі пошуку екстремумів функції "
        "за допомогою засобів Python»"
    )
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        "Самостійна практична робота з курсу «Моделі та системи штучного інтелекту». "
        "Середовище виконання: мова Python.\n\n"
        f"Усі студенти отримують матеріали з офіційного репозиторію на GitHub: {GITHUB_REPO_WEB}\n\n"
        "Версія документа: 1.5\n"
        "Формат: Python 3.10+"
    )
    mr.font.size = Pt(12)

    doc.add_page_break()

    # --- 1. Мета ---
    _heading(doc, "1. Мета практичної роботи", level=1)
    doc.add_paragraph(
        "Вивчити основні принципи генетичного алгоритму та набути практичних навичок "
        "чисельної оптимізації функцій однієї й двох змінних за допомогою генетичного "
        "алгоритму в середовищі Python (офіційний репозиторій на GitHub): запуск сценарію, аналіз результатів, "
        "побудова графіків, робота з графічним інтерфейсом."
    )

    # --- 2. Зміст завдань і виконання ---
    _heading(doc, "2. Зміст практичної роботи: завдання та реалізація у проєкті", level=1)
    doc.add_paragraph(
        "Робота складається з чотирьох основних завдань і таблиці з п’ятьма варіантами "
        "цільових функцій (розділ 7). Матеріали потрібно брати лише з офіційного GitHub-репозиторію "
        f"({GITHUB_REPO_WEB}), щоб версія коду збігалася у всіх студентів. Нижче — що зробити і який файл "
        "проєкту для цього використати."
    )
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Завдання практичної роботи"
    hdr[1].text = "Як виконати в проєкті (Python, репозиторій на GitHub)"
    rows = [
        (
            "1. Скрипт оптимізації функції згідно з обраним варіантом (таблиця варіантів).",
            "Файл main.py запускає генетичний алгоритм для варіанту --variant 1…5. "
            "Формули та межі задані у variants.py. Студент коректно запускає сценарій "
            "і заносить числові результати у звіт.",
        ),
        (
            "2. Побудувати графік функції.",
            "Команда з параметром --plot і за потреби --plot-file … — крива y(x) або "
            "поверхня z(x,y); зображення вставляється у звіт.",
        ),
        (
            "3. Оптимізувати функцію програмно, керуючи параметрами генетичного алгоритму.",
            "Через main.py викликається модуль genetic_algorithm.py: можна змінювати "
            "--population, --generations, --seed тощо. Мета — отримати розв’язок і "
            "проаналізувати вплив параметрів і випадковості на результат.",
        ),
        (
            "4. Оптимізувати функцію з використанням графічного інтерфейсу.",
            "Запуск gui.py: вікно з полями параметрів, кнопкою запуску та графіками "
            "функції й збіжності алгоритму.",
        ),
    ]
    for i, (a, b) in enumerate(rows, start=1):
        c = tbl.rows[i].cells
        c[0].text = a
        c[1].text = b

    # --- 3. Формулювання завдань ---
    _heading(doc, "3. Формулювання завдань практичної роботи", level=1)
    doc.add_paragraph(
        "Тема: дослідження генетичних алгоритмів на задачі пошуку екстремумів функції "
        "засобами Python."
    )
    doc.add_paragraph("Завдання:")
    for i, line in enumerate(
        [
            "Скласти й виконати сценарій оптимізації функції згідно з варіантом з таблиці варіантів.",
            "Побудувати графік функції.",
            "Провести оптимізацію з програмним керуванням параметрами генетичного алгоритму.",
            "Провести оптимізацію з використанням графічного інтерфейсу алгоритму.",
        ],
        start=1,
    ):
        doc.add_paragraph(f"{i}. {line}", style="List Number")
    doc.add_paragraph(
        "У таблиці варіантів (розділ 7) наведено п’ять функцій: номер рядка відповідає "
        "параметру --variant 1 … --variant 5. За потреби уточнення формул звертайтеся "
        "до викладача; код функцій можна змінити у файлі variants.py."
    )

    # --- 4. Покрокова інструкція ---
    _heading(doc, "4. Покрокова інструкція для студента (простими словами)", level=1)
    doc.add_paragraph(
        "Виконуйте кроки по порядку. Після кожного кроку занесіть у звіт те, що проситиме "
        "викладач: числа, скріншот або короткий коментар."
    )

    doc.add_paragraph("Крок 0. Отримати проєкт з GitHub і підготувати середовище", style="Heading 3")
    doc.add_paragraph(
        f"Установіть Python (3.10 або новіше) та Git (за бажанням). Завантажте код **лише** з репозиторію "
        f"{GITHUB_REPO_WEB} — або клонуванням, або кнопкою Code → Download ZIP на сайті GitHub."
    )
    doc.add_paragraph("Приклад клонування в PowerShell (папка на робочому столі):", style="List Bullet")
    doc.add_paragraph(
        f"cd $env:USERPROFILE\\Desktop\n"
        f"git clone {GITHUB_CLONE_HTTPS}\n"
        f"cd {REPO_FOLDER_NAME}",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Далі у цій самій папці виконайте встановлення залежностей (розділ 8). Переконайтеся, "
        "що команда python main.py --help виводить довідку без помилки.",
        style="List Bullet",
    )

    doc.add_paragraph("Крок 1 — завдання 1 (сценарій оптимізації за варіантом)", style="Heading 3")
    doc.add_paragraph(
        "Дізнайтеся свій номер варіанту (1…5) за списком групи або таблицею варіантів (розділ 7). "
        "За бажанням відкрийте variants.py і переконайтеся, який рядок таблиці "
        "відповідає вашому номеру (межі x та тип екстремуму: мінімум чи максимум)."
    )
    doc.add_paragraph(
        "Запустіть оптимізацію однією командою, підставивши свій номер замість N:",
        style="List Bullet",
    )
    doc.add_paragraph(
        "python main.py --variant N --generations 120 --population 40",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "У консолі з’являться «Найкраща точка» та «Значення f». Скопіюйте їх у звіт. "
        "Це і є виконання пункту «скрипт оптимізації»: вам не потрібно писати ГА з нуля — "
        "потрібно правильно запустити готовий скрипт для свого варіанту.",
        style="List Bullet",
    )

    doc.add_paragraph("Крок 2 — завдання 2 (графік функції)", style="Heading 3")
    doc.add_paragraph(
        "Побудуйте графік тієї ж функції, що й у варіанті. Збережіть малюнок у файл, "
        "щоб його легко вставити у Word:",
        style="List Bullet",
    )
    doc.add_paragraph(
        "python main.py --variant N --plot --plot-file figures\\variantN.png --no-show",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Створіть папку figures заздалегідь або змініть шлях. У звіті: вставте малюнок "
        "і 2–3 речення — що на ньому видно (ріст/спад, екстремуми, для варіантів 2 і 4 "
        "— нагадати про полюси біля x=1 та x=3).",
        style="List Bullet",
    )

    doc.add_paragraph("Крок 3 — завдання 3 (програмна оптимізація, параметри ГА)", style="Heading 3")
    doc.add_paragraph(
        "Повторіть оптимізацію з консолі, змінюючи параметри: випадковий сід, розмір "
        "популяції, кількість поколінь. Це дозволяє побачити вплив налаштувань і "
        "стохастичності генетичного алгоритму.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "python main.py --variant N --seed 1 --generations 100\n"
        "python main.py --variant N --seed 2 --generations 100\n"
        "python main.py --variant N --seed 42 --population 80 --generations 150",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Оформіть маленьку таблицю в звіті: стовпці «seed», «population», «generations», "
        "«знайдена точка», «значення f». Напишіть одним абзацом, чи змінюються результати "
        "(ГА стохастичний — це нормально).",
        style="List Bullet",
    )

    doc.add_paragraph("Крок 4 — завдання 4 (графічний інтерфейс)", style="Heading 3")
    doc.add_paragraph(
        "Запустіть графічний інтерфейс:",
        style="List Bullet",
    )
    doc.add_paragraph("python gui.py", style="Intense Quote")
    doc.add_paragraph(
        "У вікні виберіть той самий варіант N, встановіть розмір популяції та поколінь, "
        "за потреби введіть сід. Натисніть «Запустити ГА». Зробіть скріншот вікна з "
        "результатом і графіками — для розділу звіту «хід роботи».",
        style="List Bullet",
    )

    doc.add_paragraph("Крок 5. Оформлення звіту", style="Heading 3")
    doc.add_paragraph(
        "Зберіть усі скріншоти, графік і таблицю запусків. Заповніть розділи звіту "
        "згідно з розділом 10 цих методичних вказівок."
    )

    # --- 5. Структура проєкту ---
    _heading(doc, "5. З чого складається репозиторій (основні файли)", level=1)
    for item in (
        "variants.py — опис усіх п’яти варіантів (формула, межі, мінімум чи максимум);",
        "genetic_algorithm.py — сам генетичний алгоритм;",
        "plotting.py — побудова графіків;",
        "main.py — запуск з консолі (об’єднує кроки «скрипт» і «оптимізація»);",
        "gui.py — вікно з кнопкою запуску та графіками збіжності.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    # --- 6. Теорія ---
    _heading(doc, "6. Теоретичні відомості (стисло)", level=1)
    doc.add_paragraph(
        "Генетичний алгоритм — евристичний метод пошуку екстремуму: підтримується "
        "популяція рішень, кожне оцінюється цільовою функцією, кращі «виживають» "
        "частіше завдяки селекції, кросовер змішує координати батьків, мутація додає "
        "випадкові зміни. Задача на максимум у програмі зводиться до мінімізації «мінус f»."
    )
    doc.add_paragraph(
        "У реалізації: турнірний відбір, арифметичний кросовер, гаусова мутація з обрізанням "
        "по межах, елітизм; зупинка після заданої кількості поколінь (параметр generations)."
    )

    # --- 7. Таблиця варіантів ---
    _heading(doc, "7. Таблиця варіантів цільових функцій", level=1)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Функція та область"
    hdr[2].text = "Змінні"
    hdr[3].text = "Екстремум"

    rows_data = [
        ("1", "y(x) = 3x³ − 4x² + 2x − 8;  x ∈ [−6; −2]", "1", "максимум"),
        (
            "2",
            "y(x) = −6/(x−3) + 2/(x−1) + 8;  x ∈ [−4; 8]\n(штраф біля полюсів x=1, x=3)",
            "1",
            "мінімум",
        ),
        ("3", "z(x, y) = x² + y² − x·y·exp(−(x+y));  x,y ∈ [−2; 2]", "2", "мінімум"),
        ("4", "та сама y(x), що в п. 2", "1", "максимум"),
        ("5", "та сама z(x,y), що в п. 3", "2", "максимум"),
    ]
    for i, row in enumerate(rows_data, start=1):
        cells = table.rows[i].cells
        for j, txt in enumerate(row):
            cells[j].text = txt

    # --- 8. Установка ---
    _heading(doc, "8. Отримання проєкту з GitHub та установка бібліотек", level=1)
    doc.add_paragraph(
        f"Адреса репозиторію: {GITHUB_REPO_WEB}. Після клонування або розпакування ZIP "
        f"перейдіть у папку {REPO_FOLDER_NAME} і виконайте:"
    )
    doc.add_paragraph(
        f"cd $env:USERPROFILE\\Desktop\\{REPO_FOLDER_NAME}\n"
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "pip install -r requirements.txt",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "Якщо в консолі «кракозябри» замість українських літер: $env:PYTHONUTF8=1"
    )

    # --- 9. Довідник команд ---
    _heading(doc, "9. Довідник команд (коротко)", level=1)
    doc.add_paragraph(
        "python main.py --variant N [--population K] [--generations G] [--seed S] "
        "[--plot] [--plot-file шлях.png] [--no-show] [--json]",
        style="Intense Quote",
    )
    doc.add_paragraph("python gui.py", style="Intense Quote")
    doc.add_paragraph(
        "Для викладача: щоб оновити цей файл Word після правок у build_methodical_docx.py — "
        "команда python build_methodical_docx.py."
    )

    # --- 10. Зміст звіту ---
    _heading(doc, "10. Зміст звіту студента", level=1)
    doc.add_paragraph("У звіті обов’язково мають бути такі частини (у зручному для вас порядку):")
    report_items = [
        "Титульна сторінка.",
        "Мета роботи.",
        "Завдання.",
        "Скрипт-файл оптимізації функцій (опис або лістинг виклику main.py; посилання на GitHub-репозиторій).",
        "Опис виконання по пунктам завдання (хід роботи) зі скріншотами.",
        "Висновки.",
    ]
    for it in report_items:
        doc.add_paragraph(it, style="List Number")
    doc.add_paragraph(
        "У пункті 5 «хід роботи» логічно розташувати: скрін консолі після кроку 1, "
        "графік після кроку 2, таблицю експериментів після кроку 3, скрін gui.py після кроку 4."
    )

    # --- 11. Контрольні питання ---
    _heading(doc, "11. Контрольні питання та зразкові відповіді", level=1)
    qa = [
        (
            "Перерахуйте основні особливості ГА.",
            "Пошук із популяції; використовуються значення цільової функції без похідних; "
            "стохастичний відбір; ітеративне покращення наближення.",
        ),
        (
            "Перелічіть генетичні оператори.",
            "Селекція, кросовер (схрещування), мутація; часто елітизм.",
        ),
        (
            "Які критерії зупинки використовуються для ГА?",
            "Ліміт поколінь, час, поріг якості, відсутність покращення за кілька поколінь, "
            "мала відносна зміна найкращого значення функції пристосування. У цьому репозиторії "
            "реалізовано зупинку за кількістю поколінь.",
        ),
        (
            "Опишіть схему класичного ГА.",
            "Початкова популяція → оцінка пристосованості → перевірка зупинки → селекція → "
            "кросовер і мутація → нове покоління → повторення.",
        ),
        (
            "У чому полягають особливості спільного використання генетичних операторів?",
            "Кросовер поєднує вдалі частини рішень; мутація дає різноманітність і допомагає "
            "вийти з локального екстремуму; разом із селекцією це баланс між збіжністю та дослідженням простору.",
        ),
    ]
    for n, (q, a) in enumerate(qa, start=1):
        doc.add_paragraph(f"{n}. {q}", style="List Number")
        doc.add_paragraph(f"Відповідь: {a}")

    # --- 12. Додаток ---
    _heading(doc, "12. Додаток: відповідність файлів проєкту", level=1)
    files = [
        ("variants.py", "варіанти функцій, межі, min/max"),
        ("genetic_algorithm.py", "ядро ГА"),
        ("plotting.py", "графіки"),
        ("main.py", "консольний запуск"),
        ("gui.py", "графічний інтерфейс"),
        ("requirements.txt", "залежності"),
        ("build_methodical_docx.py", "генерація цього документа Word"),
    ]
    t2 = doc.add_table(rows=len(files) + 1, cols=2)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    h2[0].text = "Файл"
    h2[1].text = "Призначення"
    for i, (fn, desc) in enumerate(files, start=1):
        row = t2.rows[i].cells
        row[0].text = fn
        row[1].text = desc

    # --- 13. Література ---
    _heading(doc, "13. Рекомендована література", level=1)
    refs = [
        "Рутковська Д., Пилинський М., Рутковський Л. Нейронні мережі, генетичні алгоритми "
        "та нечіткі системи. — М.: Гаряча лінія — Телеком, 2006.",
        "Кононюк А. Ю. Нейронні мережі і генетичні алгоритми. — К.: Корнійчук, 2008.",
        "Документація Python: https://docs.python.org/3/",
    ]
    for r in refs:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.add_run("— Кінець документа —").italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(output_path))
        print("Збережено:", output_path.resolve())
    except PermissionError:
        alt = output_path.parent / "Methodical_Practice_GA_Python_UA_BUILD.docx"
        doc.save(str(alt))
        print(
            "УВАГА: основний .docx зайнятий (закрийте його в Word і запустіть скрипт знову). "
            "Збережено копію:",
            alt.resolve(),
        )


if __name__ == "__main__":
    base = Path(__file__).resolve().parent
    ua = base / "Методичні_вказівки_Практична_робота_ГА_Python.docx"
    ascii_copy = base / "Methodical_Practice_GA_Python_UA.docx"
    build(ua)
    try:
        shutil.copy2(ua, ascii_copy)
        print("Копія (ASCII ім'я):", ascii_copy.resolve())
    except PermissionError:
        print("Копію ASCII не оновлено (файл зайнятий). Використайте BUILD-файл з повідомлення вище.")
